from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mattdimock/Documents/Jobs/Job Search")
OUT_DIR = ROOT / "applications" / "hospitality-lanes"

NAVY = RGBColor(24, 53, 77)
ACCENT = RGBColor(182, 111, 55)
TEXT = RGBColor(28, 31, 36)
MUTED = RGBColor(95, 104, 113)


RESUMES = [
    {
        "filename": "Matt Dimock - Barback Bar Support.docx",
        "title": "Barback | Bar Support",
        "summary": (
            "Reliable, physically capable operations professional with recent experience managing fast-moving work, tight handoffs, and high expectations, "
            "plus earlier retail team-lead experience in stocking, customer service, opening and closing, and department upkeep. "
            "Familiar with common spirits and classic cocktail structure from years of home practice, and ready to earn trust quickly in a high-volume bar environment."
        ),
        "strengths": [
            "Immediate availability for nights, weekends, and holidays",
            "Physically comfortable with long shifts, heavy lifting, and constant movement",
            "Strong stocking, reset, cleanliness, and side-work discipline",
            "Calm under pressure with clear communication and follow-through",
            "Fast learner with strong memory for house standards and drink builds",
            "Team-first mindset and low ego in support roles",
        ],
        "experience": [
            (
                "Get Fractional",
                "Growth Systems Consultant / Owner",
                "Oct 2022-Present",
                [
                    "Managed fast-moving client and operator work across multiple priorities, built repeatable checklists and workflows, and kept execution clean under pressure.",
                ],
            ),
            (
                "Prosper Wireless",
                "Director of Growth & Retention",
                "Sept 2023-Nov 2025",
                [
                    "Supported a large customer operation through training, documentation, and process discipline, strengthening responsiveness and consistency in a high-volume environment.",
                ],
            ),
            (
                "Bob's Watches",
                "Director of Marketing",
                "Jan 2017-Dec 2018",
                [
                    "Coordinated cross-functional execution across staff and vendors in a busy ecommerce business where timing, standards, and customer-impacting details mattered every day.",
                ],
            ),
            (
                "Lowe's Home Improvement",
                "Team Leader, Paint Department",
                "Mar 2005-Feb 2007",
                [
                    "Promoted into a team-lead role within six months and oversaw customer service, merchandising, stocking, and day-to-day department execution.",
                    "Opened and closed the department, supported a four-person team, and kept inventory, shelves, and customer flow organized during busy periods.",
                ],
            ),
            (
                "Boething Treeland Nursery",
                "Guide",
                "Dec 2003-Aug 2004",
                [
                    "Promoted from Loader to Guide within three months in a wholesale and retail nursery.",
                    "Loaded customer vehicles, restocked product areas, maintained grounds, and helped customers find the right products quickly.",
                ],
            ),
        ],
        "additional": [
            "Local to Mount Juliet and open to downtown Nashville work",
            "Familiar with common spirits, mixers, garnishes, and classic cocktail structure",
            "Ready to complete TABC-certified alcohol awareness training immediately",
            "Interested in growing into bartending after proving value in support work",
        ],
    },
    {
        "filename": "Matt Dimock - Food Runner.docx",
        "title": "Food Runner",
        "summary": (
            "Fast-moving, customer-facing operations professional with recent experience handling multiple priorities, tight handoffs, and high service expectations, "
            "backed by earlier retail leadership in stocking, customer support, and keeping busy floor environments organized. "
            "Brings urgency, accuracy, and a team-first mindset to high-volume food runner work."
        ),
        "strengths": [
            "Immediate availability for nights, weekends, and holidays",
            "Strong urgency, timing, and order accuracy",
            "Clear communication with kitchen and floor teams",
            "Comfortable on feet for long shifts and constant movement",
            "Detail-focused with strong follow-through",
            "Guest-friendly and professional under pressure",
        ],
        "experience": [
            (
                "Get Fractional",
                "Growth Systems Consultant / Owner",
                "Oct 2022-Present",
                [
                    "Managed high-priority work with shifting timelines, multiple handoffs, and demanding expectations while keeping communication clear and details accurate.",
                ],
            ),
            (
                "Breakthrough Academy",
                "Marketing Operations Manager",
                "Aug 2019-Jun 2021",
                [
                    "Coordinated events, webinars, partners, and vendors where timing, presentation, and clean execution directly affected the customer experience.",
                ],
            ),
            (
                "Bob's Watches",
                "Director of Marketing",
                "Jan 2017-Dec 2018",
                [
                    "Helped lead execution across teams in a fast-paced ecommerce business where accuracy, responsiveness, and customer-impacting details mattered.",
                ],
            ),
            (
                "Lowe's Home Improvement",
                "Team Leader, Paint Department",
                "Mar 2005-Feb 2007",
                [
                    "Oversaw customer service, merchandising, stocking, and department flow in a busy retail environment.",
                    "Managed a four-person team and helped customers make confident purchase decisions quickly and professionally.",
                ],
            ),
            (
                "Boething Treeland Nursery",
                "Guide",
                "Dec 2003-Aug 2004",
                [
                    "Worked in a fast-moving retail setting where physical pace, product familiarity, and customer service mattered.",
                    "Loaded purchases, maintained inventory areas, and helped customers navigate options efficiently.",
                ],
            ),
        ],
        "additional": [
            "Local to Mount Juliet and open to downtown Nashville work",
            "Physically active and comfortable with high-volume support work",
            "Strong memory and quick study habits for menus, table numbers, and seat positions",
            "Ready to start immediately",
        ],
    },
    {
        "filename": "Matt Dimock - Server Assistant.docx",
        "title": "Server Assistant",
        "summary": (
            "Customer-facing operations professional with a strong record of staying organized, supporting teams, and handling multiple priorities without losing detail. "
            "Recent work reflects high standards, strong communication, and clean execution under pressure, while earlier retail leadership adds direct proof in floor support, stocking, guest interaction, and team coordination."
        ),
        "strengths": [
            "Immediate availability for nights, weekends, and holidays",
            "Strong pace, awareness, and follow-through",
            "Professional with guests and polished under pressure",
            "Strong reset, side-work, and station-discipline habits",
            "Comfortable supporting busy teams without constant direction",
            "Fast learner with dependable shift discipline",
        ],
        "experience": [
            (
                "Get Fractional",
                "Growth Systems Consultant / Owner",
                "Oct 2022-Present",
                [
                    "Managed demanding work across multiple stakeholders, kept priorities organized, and maintained strong follow-through in high-expectation environments.",
                ],
            ),
            (
                "Prosper Wireless",
                "Director of Growth & Retention",
                "Sept 2023-Nov 2025",
                [
                    "Built training and process systems that improved consistency, helped teams stay coordinated, and supported a large-volume customer operation.",
                ],
            ),
            (
                "Breakthrough Academy",
                "Marketing Operations Manager",
                "Aug 2019-Jun 2021",
                [
                    "Coordinated events, partners, and internal teams where presentation, timing, and responsiveness directly affected the customer experience.",
                ],
            ),
            (
                "Lowe's Home Improvement",
                "Team Leader, Paint Department",
                "Mar 2005-Feb 2007",
                [
                    "Oversaw customer service, sales, merchandising, and day-to-day department execution in a busy retail setting.",
                    "Opened and closed the department, supported a four-person team, and helped keep standards high across stocking, service, and floor organization.",
                ],
            ),
            (
                "Boething Treeland Nursery",
                "Guide",
                "Dec 2003-Aug 2004",
                [
                    "Helped customers find the right products, maintained product areas, and worked in a physical, customer-facing environment where pace and follow-through mattered.",
                ],
            ),
        ],
        "additional": [
            "Local to Mount Juliet and open to downtown Nashville work",
            "Open to high-volume shifts and strong support-floor expectations",
            "Strong at anticipating needs, staying composed, and protecting service flow",
            "Ready to start immediately",
        ],
    },
    {
        "filename": "Matt Dimock - Host Door Greeter.docx",
        "title": "Host | Door Greeter",
        "summary": (
            "Calm, professional, customer-facing team member with recent experience handling high expectations, difficult personalities, and moving priorities without losing composure. "
            "Earlier retail and team-lead experience adds direct proof in customer service, floor awareness, team coordination, and maintaining standards in busy public-facing environments."
        ),
        "strengths": [
            "Immediate availability for nights, weekends, and holidays",
            "Strong communication and calm presence",
            "Comfortable with crowds and high expectations",
            "Professional problem-solving and issue routing",
            "Dependable, punctual, and organized",
            "Good judgment under pressure",
        ],
        "experience": [
            (
                "Get Fractional",
                "Growth Systems Consultant / Owner",
                "Oct 2022-Present",
                [
                    "Managed demanding client relationships where responsiveness, tone, and professionalism mattered at all times.",
                ],
            ),
            (
                "Prosper Wireless",
                "Director of Growth & Retention",
                "Sept 2023-Nov 2025",
                [
                    "Supported training, communication, and process quality across a large customer operation where clarity and consistency mattered.",
                ],
            ),
            (
                "Breakthrough Academy",
                "Marketing Operations Manager",
                "Aug 2019-Jun 2021",
                [
                    "Helped coordinate member events, webinars, and partner experiences where presentation, organization, and guest experience were important.",
                ],
            ),
            (
                "Lowe's Home Improvement",
                "Team Leader, Paint Department",
                "Mar 2005-Feb 2007",
                [
                    "Worked directly with customers, handled product questions, and helped people make confident decisions in a busy retail setting.",
                    "Managed a four-person team and maintained service and presentation standards across the department.",
                ],
            ),
            (
                "Boething Treeland Nursery",
                "Guide",
                "Dec 2003-Aug 2004",
                [
                    "Guided customers through product choices, helped load purchases, and supported a smooth customer experience from entry through checkout.",
                ],
            ),
        ],
        "additional": [
            "Local to Mount Juliet and open to downtown Nashville work",
            "Comfortable with high-volume guest interaction and strong front-door expectations",
            "Strong memory for names, details, and repeat interactions",
            "Ready to start immediately",
        ],
    },
]


def set_margins(section):
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)


def set_bottom_border(paragraph, color="B66F37", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_header(doc, spec):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("Matt Dimock")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(spec["title"])
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Mount Juliet, TN  |  805-620-2826  |  mattdim805@gmail.com  |  linkedin.com/in/mattdimock/")
    run.font.name = "Aptos"
    run.font.size = Pt(9.2)
    run.font.color.rgb = MUTED
    set_bottom_border(p)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY


def add_body_paragraph(doc, text, font_size=9.4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(font_size)
    run.font.color.rgb = TEXT


def add_bullets(doc, items, font_size=9.2):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.first_line_indent = Pt(-8)
        run = p.add_run(f"• {item}")
        run.font.name = "Aptos"
        run.font.size = Pt(font_size)
        run.font.color.rgb = TEXT


def add_experience(doc, company, title, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(company)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT
    run = p.add_run(f"  |  {title}  |  {dates}")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    add_bullets(doc, bullets, font_size=9.0)


def build_resume(spec):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    set_margins(section)

    add_header(doc, spec)
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, spec["summary"])
    add_section_heading(doc, "Core Strengths")
    add_bullets(doc, spec["strengths"])
    add_section_heading(doc, "Relevant Experience")
    for company, title, dates, bullets in spec["experience"]:
        add_experience(doc, company, title, dates, bullets)
    add_section_heading(doc, "Additional Information")
    add_bullets(doc, spec["additional"])

    out_path = OUT_DIR / spec["filename"]
    doc.save(out_path)
    return out_path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for spec in RESUMES:
        build_resume(spec)
        print(f"rendered {spec['filename']}")


if __name__ == "__main__":
    main()
